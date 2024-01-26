# word_tokenizer.py
import nltk
nltk.download('punkt')

from nltk.tokenize import word_tokenize
from docx import Document

def tokenize_text(text):
    return word_tokenize(text)

def tokenize_word_document(docx_path):
    doc = Document(docx_path)
    tokens = []
    for paragraph in doc.paragraphs:
        tokens.extend(word_tokenize(paragraph.text))
    return tokens

def replace_tokens(docx_path, replacement_dict):
    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for token, replacement in replacement_dict.items():
                text = text.replace(token, replacement)
            run.text = text
    return doc
