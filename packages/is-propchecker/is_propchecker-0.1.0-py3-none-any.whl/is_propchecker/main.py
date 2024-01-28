import typer
import docx
from docx import Document
import os
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH , WD_LINE_SPACING , WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from PyPDF2 import PdfReader
import docx2pdf


app = typer.Typer()

SBIR_Ag = ['navy', 'airforce', 'mda', 'dha', 'dla', 'dmea']
STTR_Ag = ['Airforce', 'navy', 'dmea']

@app.command()
def hello(name1: str, name2: str, path, phase : int, proposal_num, topic_name, comp):
    '''
    if name1 == 'SBIR':
        if name2 == 'navy':
            page_height = 27.94
            page_width = 21.59
            left_margin = 2.54
            right_margin = 2.54
            top_margin = 2.54
            bottom_margin = 2.54
            header_distance = 1.27
            footer_distance = 1.27
            font_name = 'Times New Roman'
            font_size = 10

        if name2 == 'airforce':
            page_height = 27.94
            page_width = 21.59
            left_margin = 2.54
            right_margin = 2.54
            top_margin = 2.54
            bottom_margin = 2.54
            header_distance = 1.27
            footer_distance = 1.27
            font_name = 'Times New Roman'
            font_size = 10

        if name2 == 'mda':
            page_height = 27.94
            page_width = 21.59
            left_margin = 2.54
            right_margin = 2.54
            top_margin = 2.54
            bottom_margin = 2.54
            header_distance = 1.27
            footer_distance = 1.27
            font_name = 'Times New Roman'
            font_size = 10

        if name2 == 'dha':
            page_height = 27.94
            page_width = 21.59
            left_margin = 2.54
            right_margin = 2.54
            top_margin = 2.54
            bottom_margin = 2.54
            header_distance = 1.27
            footer_distance = 1.27
            font_name = 'Times New Roman'
            font_size = 10

        if name2 == 'dla':
            page_height = 27.94
            page_width = 21.59
            left_margin = 2.54
            right_margin = 2.54
            top_margin = 2.54
            bottom_margin = 2.54
            header_distance = 1.27
            footer_distance = 1.27
            font_name = 'Arial'
            font_size = 10

        if name2 == 'dmea':
            page_height = 27.94
            page_width = 21.59
            left_margin = 2.54
            right_margin = 2.54
            top_margin = 2.54
            bottom_margin = 2.54
            header_distance = 1.27
            footer_distance = 1.27
            font_name = 'Arial'
            font_size = 10
    
    elif name1 == 'STTR':
        if name2 == 'navy':
            page_height = 27.94
            page_width = 21.59
            left_margin = 2.54
            right_margin = 2.54
            top_margin = 2.54
            bottom_margin = 2.54
            header_distance = 1.27
            footer_distance = 1.27
            font_name = 'Arial'
            font_size = 10

        if name2 == 'Airforce':
            page_height = 27.94
            page_width = 21.59
            left_margin = 2.54
            right_margin = 2.54
            top_margin = 2.54
            bottom_margin = 2.54
            header_distance = 1.27
            footer_distance = 1.27
            font_name = 'Arial'
            font_size = 10

        if name2 == 'dmea':
            page_height = 27.94
            page_width = 21.59
            left_margin = 2.54
            right_margin = 2.54
            top_margin = 2.54
            bottom_margin = 2.54
            header_distance = 1.27
            footer_distance = 1.27
            font_name = 'Arial'
            font_size = 10        
    else:
        print('Please choose between SBIR and STTR')
    '''
    page_height = 27.94
    page_width = 21.59
    left_margin = 2.54
    right_margin = 2.54
    top_margin = 2.54
    bottom_margin = 2.54
    header_distance = 1.27
    footer_distance = 1.27
    font_name = 'Times New Roman'
    font_size = 10

    if (name1 =='STTR' or 'SBIR') and (name2 == 'navy' or 'airforce'):
        if phase == 1:
            page_limit = 10
            print('1')
        else:
            page_limit = 30
            print('2')

    elif (name1 =='STTR' or 'SBIR') and (name2 == 'dmea'):
        page_limit = 20
        print('3')
    elif (name1 =='SBIR') and (name2 == 'dha' or 'dla'):
        if phase == 1:
            page_limit = 20
            print('4')
        else:
            page_limit = 60
            print('4')
    elif (name1 =='STTR' or 'SBIR') and (name2 == 'mda'):
        page_limit = 15
        print('5')
    else:
        print("check your input")

    _doc = path

    # print('hello done')
    print(comp, topic_name, proposal_num, _doc, page_height, page_width,left_margin,right_margin, top_margin,bottom_margin,header_distance,footer_distance,font_name,font_size)
    print(phase, name1, name2, page_limit)
    format(page_limit, comp, topic_name, proposal_num, _doc, page_height, page_width,left_margin,right_margin, top_margin,bottom_margin,header_distance,footer_distance,font_name,font_size)
    
    return comp, topic_name, proposal_num, _doc, page_height, page_width,left_margin,right_margin, top_margin,bottom_margin,header_distance,footer_distance,font_name,font_size


def format(page_limit, comp, topic_name, proposal_num, _doc, page_height, page_width,left_margin,right_margin, top_margin,bottom_margin,header_distance,footer_distance,font_name,font_size):

    _doc = _doc
    print('proposal path : ', _doc)

    document = Document(_doc)
    # print(document.paragraphs)

    section = document.sections[0]
    section.page_height = Cm(page_height) # Defining the paper size to US Letter
    section.page_width = Cm(page_width)
    section.left_margin = Cm(left_margin) # Defining the side Margines of the document
    section.right_margin = Cm(right_margin)
    section.top_margin = Cm(top_margin)
    section.bottom_margin = Cm(bottom_margin)
    section.header_distance = Cm(header_distance) # Defining the Header distance
    section.footer_distance = Cm(footer_distance) # Defining the footer distance
    
    n = str(proposal_num)
    m = str(topic_name)
    p = str(comp)
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = f'Proposal Number - {n}\t{p}\nTopic Name - {m}'
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    for p in document.paragraphs:
        # print(p.text) #print the paragraphs in new lines
        
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY #justify the lines for each paragraph
        p.paragraph_format.line_spacing = Pt(0) #Lines spacing in the content of the paragraphs
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY # Defining the rule of the Line spacing that the content of the paragraphs will have
        
        for run in p.runs:
            run.font.name = str(font_name) #font name of the content of the desired paragraph
            run.font.size = Pt(font_size) #font size of the content of the desired paragraph


    # doc = Document(_doc)

    add_page_number(document.sections[0].footer.paragraphs[0].add_run())
    document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.sections[0].different_first_page_header_footer = False
    sectPr = document.sections[0]._sectPr

    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(ns.qn('w:start'), "1")
    sectPr.append(pgNumType)
    # print(sectPr)

    document.save(str(_doc[:-5]+'_result.docx')) #Path for saving the edited file
    
    if not os.path.exists('./pdfs'):
        os.mkdir('./pdfs') #Destination folder path for working on Document

    d_doc = str(_doc[:-5]+'_result.docx')
    d__doc = os.path.join('./pdfs',str(str(n)+'.pdf'))

    p = docx2pdf.convert(d_doc,d__doc) # COnvertion of the Documnet file to PDF file
    r = PdfReader(str(d__doc)) #PDFReader can read the pdf and extract information we like.
    num_pages = len(r.pages) #Counting the number of pages 

    if num_pages > page_limit:
        print(f"You should reduce the content by {num_pages - page_limit} page(s)")
    else:
        print(f"It is short by {page_limit - num_pages} page(s)")

    return str(_doc[:-5]+'_result.docx')
    
def create_element(name):
    return OxmlElement(name)
        
def create_attribute(element, name, value):
    element.set(ns.qn(name), value)
    # print(vlaue)
        
def add_page_number(run):
    fldStart = create_element('w:fldChar')
    create_attribute(fldStart, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'separate')

    fldChar2 = create_element('w:t')
    fldChar2.text = "2"

    fldEnd = create_element('w:fldChar')
    create_attribute(fldEnd, 'w:fldCharType', 'end')

    run._r.append(fldStart)

    run._r.append(instrText)
    run._r.append(fldChar1)
    run._r.append(fldChar2)

    run._r.append(fldEnd)
    # print('hi ', run)

# if __name__ =="__main__":
#     app()